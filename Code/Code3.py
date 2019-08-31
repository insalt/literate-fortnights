# your code goes here! /usr/bin/python
from nltk.tokenize import sent_tokenize
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_UNDERLINE
import glob
import re

for filename in glob.iglob('/definition-doc/*.docx'):
    print ("x")
    print('/definition-doc/%s' % filename)
    fullText = []
    for para in contract.paragraphs:
        fullText.append(para.text)
    x = sent_tokenize("\n".join(fullText))



def Count_Total_2():
    with open("New_Dict.txt") as New_Dict:
        pattern = re.compile(r'\b(' + '|'.join(New_Dict.keys()) + r')\b')
        result = pattern.sub(lambda z: d[z.group()], x)
        print (result)
    # return (Broad_Terms_Found)

Count_Total_2()

# def Count_Broad():
#     with open("Broad_Terms.txt") as Broad_Dict:
#         Broad_Dict = [y.strip() for y in Broad_Dict.readlines()]
#     total_positive_score = 0
#     n = 0
#     contract.paragraphs[-1].add_run("\n\nBroad Terms Found").underline = WD_UNDERLINE.SINGLE
#     while n < len(x):
#         p=1
#         for i in Broad_Dict:
#                 if i in x[n]:
#                     total_positive_score += p
#                     p+=1
#                     # global Broad_Terms_Found
#                     # Broad_Terms_Found.join("{}".format(i))
#                     contract.paragraphs[-1].add_run("\n"+(i)).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
#                     # contract.save('test2.docx')
#         n+=1
#     return (total_positive_score)
#     # return (Broad_Terms_Found)
#
# def Count_Spec():
#     with open("Specific_Terms.txt") as Spec_Dict:
#         Spec_Dict = [z.strip() for z in Spec_Dict.readlines()]
#     total_neg_score = 0
#     m = 0
#     contract.paragraphs[-1].add_run("\n\nSpecific Terms Found").underline = WD_UNDERLINE.SINGLE
#     while m < len(x):
#         q=1
#         for i in Spec_Dict:
#             if i in x[m]:
#                 total_neg_score = total_neg_score - q
#                 q=q+1
#                 # global Spec_Terms_Found
#                 # Spec_Terms_Found.join(i)
#                 contract.paragraphs[-1].add_run("\n"+(i)).font.highlight_color = WD_COLOR_INDEX.RED
#                 # contract.save('test2.docx')
#         m+=1
#     return total_neg_score
    # return Spec_Terms_Found

# def Count_Total(a,b):
#     x = abs(a) + abs(b)
#     y = round(abs(((a+b)/x)*100), 1)
#     # if a+b == 0:
#         # contract.paragraphs[0].add_run("This clause is a netrally phrased clause.")
#         # contract.paragraphs[0].add_run("SUMMARY")
#         # contract.paragraphs[0].add_run("Broad Terms Found: "+ "".join(Broad_Terms_Found).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
#         # contract.paragraphs[0].add_run("Specific Terms Found: "+ "".join(Spec_Terms_Found).font.highlight_color = WD_COLOR_INDEX.RED
#         # contract.save('test.docx')
#     if a+b<0:
#         contract.paragraphs[-1].add_run("\n\nThis clause is {}% phrased in the counter party's favour.".format(y))
#         # contract.paragraphs[0].add_run("\nBroad Terms Found: "+ Broad_Terms_Found)
#         # contract.paragraphs[0].add_run("\nSpecific Terms Found: "+ Spec_Terms_Found)
#         contract.save('definition-doc/PartyA/%s' %s filename )
#     elif a+b >0:
#         contract.paragraphs[-1].add_run("\n\nThis clause is {}% phrased in the counter party's favour.".format(y))
#         # contract.paragraphs[0].add_run("\nBroad Terms Found: "+ Broad_Terms_Found)
#         # contract.paragraphs[0].add_run("\nSpecific Terms Found: "+ Spec_Terms_Found)
#         contract.save('definition-doc/PartyB/%s' %s filename )
