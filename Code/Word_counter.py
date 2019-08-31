import nltk
from nltk.tokenize import word_tokenize
import numpy as np
import random
import pickle
from collections import Counter
from nltk.stem import WordNetLemmatizer
from nltk import FreqDist
import glob
import docx

def create_word_counter():
	features = []
	for filename in glob.iglob('PartA/*.docx'):
			contract = docx.Document('%s.docx' %filename)
			fullText = []
			for para in contract.paragraphs:
				fullText.append(para.text)
				x = sent_tokenize("\n".join(fullText))
			features += x
			print ("x")
	for filename in glob.iglob('PartB/*.docx'):
			contract = docx.Document('%s.docx' %filename)
			fullText2 = []
			for para in contract.paragraphs:
				fullText2.append(para.text)
				y = sent_tokenize("\n".join(fullText2))
			features += y
	# fdist1 = FreqDist(features)
	# print (features)
	# print (fdist1.most_common(50))


create_word_counter()
